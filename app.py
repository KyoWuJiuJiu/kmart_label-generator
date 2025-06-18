import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import RGBColor, Pt
from io import BytesIO
from datetime import date
import os
import copy
import re

FILL_MAP = {
    (2, 4): "Assortment Breakdown",
    (3, 4): ("FOB Point", "FOB NB"),
    (4, 4): "=today()",
    (5, 1): "ITEM#",
    (8, 2): "Item Description"
}

def fill_label_table(table, data_row):
    for (ri, ci), source in FILL_MAP.items():
        if ri >= len(table.rows):
            continue
        row = table.rows[ri]
        if ci >= len(row.cells):
            continue
        target_cell = row.cells[ci]

        if source == "=today()":
            value = str(date.today())
        elif isinstance(source, tuple):
            values = [str(data_row.get(col, "")) for col in source]
            value = " / ".join(values)
        else:
            value = str(data_row.get(source, ""))

        if target_cell.paragraphs:
            para = target_cell.paragraphs[0]
            run = para.runs[0] if para.runs else para.add_run()
            run.text = value
        else:
            para = target_cell.add_paragraph()
            run = para.add_run(value)

def duplicate_table_to_new_section(doc, table):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    new_table = copy.deepcopy(table)
    doc._body._element.append(copy.deepcopy(OxmlElement('w:br')))
    doc._body._element.append(new_table._element)
    return new_table

st.title("🌍 标贴填充自动化工具")

uploaded_excel = st.file_uploader("上传Excel数据", type=["xlsx"])

if uploaded_excel:
    df = pd.read_excel(uploaded_excel)
    st.success(f"成功读取 {len(df)} 条数据")

    required_cols = set()
    for v in FILL_MAP.values():
        if isinstance(v, str) and v and not v.startswith("="):
            required_cols.add(v)
        elif isinstance(v, tuple):
            required_cols.update(v)

    missing_cols = [col for col in required_cols if col not in df.columns]
    if missing_cols:
        st.error("Excel 缺少必要列：" + ", ".join(missing_cols))
        st.stop()

    TEMPLATE_PATH = "Kmart Buy Trip Label Template.docx"
    if not os.path.exists(TEMPLATE_PATH):
        st.error("未找到固定模板文件")
    else:
        doc = Document(TEMPLATE_PATH)
        big_table_template = doc.tables[0]

        template_table = None
        for row in big_table_template.rows:
            for i, cell in enumerate(row.cells):
                if i == 1:
                    continue
                if cell.tables:
                    template_table = cell.tables[0]
                    break
            if template_table:
                break

        if not template_table:
            st.error("未找到标贴模板")
            st.stop()

        all_label_tables = []
        rows_needed = len(df)
        labels_per_table = sum(1 for row in big_table_template.rows for i in [0,2] if len(row.cells) > i)
        num_full_tables = (rows_needed + labels_per_table - 1) // labels_per_table

        for t in range(num_full_tables):
            new_big_table = copy.deepcopy(big_table_template)
            doc._body._element.append(new_big_table._element)

            for row in new_big_table.rows:
                for i, cell in enumerate(row.cells):
                    if i == 1:
                        continue
                    cell._element.clear_content()
                    new_table = copy.deepcopy(template_table)
                    cell._element.append(new_table._element)
                    all_label_tables.append(new_table)

        st.info(f"总输出 {len(all_label_tables)} 个标贴区域")

        preview_index = st.number_input("预览标贴索引", min_value=0, max_value=len(all_label_tables)-1, value=0)

        if st.button("📃 预览指定标贴"):
            fill_label_table(all_label_tables[preview_index], df.iloc[preview_index % len(df)])
            preview_output = BytesIO()
            doc.save(preview_output)
            st.download_button("🔳 下载预览标贴", data=preview_output.getvalue(), file_name=f"Preview_Label_{preview_index}.docx")

        if st.button("🚀 填充所有标贴"):
            for i, (_, row) in enumerate(df.iterrows()):
                if i >= len(all_label_tables):
                    break
                fill_label_table(all_label_tables[i], row)

            output = BytesIO()
            doc.save(output)
            st.download_button("🔳 下载生成文档", data=output.getvalue(), file_name="Filled_Labels.docx")
